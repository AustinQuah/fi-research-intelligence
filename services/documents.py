from pathlib import Path
import re


def parse_uploaded_document(
    path: str,
    filename: str,
) -> dict:

    suffix = Path(
        filename
    ).suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(path)

    if suffix == ".docx":
        return parse_docx(path)

    if suffix in {".txt", ".md"}:

        return {
            "filename": filename,
            "text": Path(path).read_text(
                encoding="utf-8",
                errors="ignore",
            ),
            "pages": None,
            "visual_pages": [],
        }

    raise ValueError(
        "Supported formats: PDF, DOCX, TXT and MD."
    )


def parse_pdf(path: str) -> dict:

    import pymupdf

    document = pymupdf.open(path)

    pages = []
    visual_pages = []

    for number, page in enumerate(
        document,
        start=1,
    ):

        text = page.get_text(
            "text"
        ) or ""

        pages.append(
            f"[PAGE {number}]\n{text}"
        )

        # Fast pass:
        # only flag sparse pages.
        # We do not OCR here.
        if len(text.strip()) < 250:

            visual_pages.append(
                number
            )

    page_count = len(
        document
    )

    document.close()

    return {
        "filename": Path(path).name,
        "text": "\n\n".join(
            pages
        ),
        "pages": page_count,
        "visual_pages": visual_pages,
    }


def parse_docx(path: str) -> dict:

    from docx import Document

    document = Document(path)

    parts = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:
            parts.append(text)

    for table in document.tables:

        for row in table.rows:

            parts.append(
                " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                )
            )

    return {
        "filename": Path(path).name,
        "text": "\n".join(parts),
        "pages": None,
        "visual_pages": [],
    }


def quick_understanding(
    parsed: dict,
) -> dict:

    text = parsed.get(
        "text",
        ""
    )

    lower = text.lower()

    # --------------------------------------------------------
    # Title
    # --------------------------------------------------------

    title = None

    pattern = (
        r"(?:title of research project|"
        r"research project title|"
        r"proposal title|"
        r"project title)"
        r"\s*[:\-]?\s*([^\n]{8,180})"
    )

    match = re.search(
        pattern,
        text,
        re.IGNORECASE,
    )

    if match:

        title = match.group(
            1
        ).strip()

    # --------------------------------------------------------
    # FI detection
    # --------------------------------------------------------

    if (
        "living lab" in lower
        and "water" in lower
    ):

        fi = "Living Lab (Water)"

    elif (
        "industrial water solutions"
        in lower
        or "wafer fab" in lower
    ):

        fi = (
            "Industrial Water Solutions (IWS)"
        )

    elif (
        "municipal water" in lower
        or "mwtd" in lower
    ):

        fi = (
            "Municipal Water: "
            "Technology Development (MWTD)"
        )

    elif (
        "competitive funding for water research"
        in lower
    ):

        fi = (
            "Competitive Funding "
            "for Water Research"
        )

    else:

        fi = None

    # --------------------------------------------------------
    # Document type
    # --------------------------------------------------------

    if (
        "funding initiative" in lower
        and "desired outcomes" in lower
    ):

        document_type = (
            "FI / programme paper"
        )

    elif (
        "project proposal" in lower
        or "scientific abstract" in lower
    ):

        document_type = (
            "Individual R&D proposal"
        )

    else:

        document_type = "Unknown"

    # --------------------------------------------------------
    # Section detection
    # --------------------------------------------------------

    section_names = [

        "Scientific Abstract",
        "Problem Statement",
        "Research Objectives",
        "Technical KPIs",
        "Methodology",
        "Landscape Scan",
        "Innovativeness",
        "Commercialisation",
        "Milestones",
        "Budget",
        "Impact",
        "TRL",

    ]

    sections = []

    for name in section_names:

        if name.lower() in lower:

            sections.append({

                "name":
                    name,

                "confidence":
                    0.50,

            })

    # --------------------------------------------------------
    # Potential claims
    # --------------------------------------------------------

    claims = []

    for line in text.splitlines():

        line = line.strip()

        if len(line) < 40:
            continue

        if re.search(

            r"novel|innovative|"
            r"improv|increase|"
            r"reduce|demonstrat|"
            r"achiev|target|"
            r"performance",

            line,

            re.IGNORECASE,

        ):

            claims.append(
                line
            )

        if len(claims) >= 10:
            break

    # --------------------------------------------------------
    # Simple summary
    # --------------------------------------------------------

    meaningful_lines = [

        x.strip()

        for x in text.splitlines()

        if len(
            x.strip()
        ) > 80

    ]

    summary = None

    if meaningful_lines:

        summary = " ".join(
            meaningful_lines[:3]
        )

        if len(summary) > 700:

            summary = (
                summary[:700]
                + "..."
            )

    return {

        "document": {

            "filename":
                parsed.get(
                    "filename"
                ),

            "title":
                title,

            "funding_initiative":
                fi,

            "document_type":
                document_type,

            "pages":
                parsed.get(
                    "pages"
                ),

            "visual_pages":
                parsed.get(
                    "visual_pages",
                    []
                ),

            "sections":
                sections,

        },

        "understanding": {

            "problem":
                None,

            "technology":
                None,

            "baseline":
                None,

            "proposed_solution":
                None,

            "novelty_claims":
                [],

            "trl_start":
                None,

            "trl_target":
                None,

            "commercialisation":
                None,

            "prior_projects":
                [],

        },

        "claims":
            claims,

        "kpis":
            [],

        "summary":
            summary,

        "review_flags": [

            {

                "severity":
                    "Review",

                "title":
                    "Fast-pass interpretation",

                "detail":
                    (
                        "This is a rapid structural "
                        "pass. Run the AI layer for "
                        "semantic interpretation."
                    ),

            }

        ],

        "raw_text":
            text,

    }
