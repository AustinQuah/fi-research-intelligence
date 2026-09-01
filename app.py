import os
import asyncio
import tempfile
from pathlib import Path
from urllib.parse import quote

import httpx
from nicegui import ui, run


APP_TITLE = "FI Research Intelligence"


# ============================================================
# COLOURS
# ============================================================

ui.colors(
    primary="#0EA5E9",
    secondary="#38BDF8",
    accent="#0284C7",
    positive="#10B981",
    warning="#F59E0B",
    negative="#EF4444",
)


# ============================================================
# LIGHT CUSTOM STYLING
# ============================================================

ui.add_head_html(
    """
    <style>

    body {
        background: #f5fbff;
        color: #0f172a;
    }

    .fi-shell {
        max-width: 1200px;
        margin: 0 auto;
    }

    .fi-card {
        border: 1px solid #dceff8 !important;
        border-radius: 18px !important;
        box-shadow: none !important;
    }

    .fi-card:hover {
        border-color: #b9dff0 !important;
    }

    </style>
    """
)


# ============================================================
# DOCUMENT PARSING
# ============================================================

def parse_pdf(path: str) -> dict:

    import pymupdf

    document = pymupdf.open(path)

    pages = []
    visual_pages = []

    for page_number, page in enumerate(
        document,
        start=1,
    ):

        text = page.get_text(
            "text"
        ) or ""

        pages.append(
            f"[PAGE {page_number}]\n{text}"
        )

        # Fast first pass.
        # Do not OCR or inspect every image yet.
        if len(text.strip()) < 250:

            visual_pages.append(
                page_number
            )

    page_count = len(
        document
    )

    document.close()

    return {

        "text":
            "\n\n".join(pages),

        "pages":
            page_count,

        "visual_pages":
            visual_pages,

    }


def parse_docx(path: str) -> dict:

    from docx import Document

    document = Document(
        path
    )

    parts = []

    for paragraph in document.paragraphs:

        text = paragraph.text.strip()

        if text:

            parts.append(
                text
            )

    for table in document.tables:

        for row in table.rows:

            parts.append(

                " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                )

            )

    return {

        "text":
            "\n".join(parts),

        "pages":
            None,

        "visual_pages":
            [],

    }


def parse_text(path: str) -> dict:

    return {

        "text":
            Path(path).read_text(
                encoding="utf-8",
                errors="ignore",
            ),

        "pages":
            None,

        "visual_pages":
            [],

    }


def parse_document(
    path: str,
    filename: str,
) -> dict:

    suffix = Path(
        filename
    ).suffix.lower()

    if suffix == ".pdf":

        return parse_pdf(
            path
        )

    if suffix == ".docx":

        return parse_docx(
            path
        )

    if suffix in {
        ".txt",
        ".md",
    }:

        return parse_text(
            path
        )

    raise ValueError(
        "Supported formats: PDF, DOCX, TXT, MD"
    )


# ============================================================
# FAST PROPOSAL UNDERSTANDING
# ============================================================

def understand_document(
    parsed: dict,
) -> dict:

    import re

    text = parsed.get(
        "text",
        ""
    )

    lower = text.lower()

    # --------------------------------------------------------
    # TITLE
    # --------------------------------------------------------

    title = None

    title_pattern = (

        r"(?:title of research project|"
        r"research project title|"
        r"proposal title|"
        r"project title)"

        r"\s*[:\-]?\s*"

        r"([^\n]{8,180})"

    )

    match = re.search(
        title_pattern,
        text,
        re.IGNORECASE,
    )

    if match:

        title = (
            match
            .group(1)
            .strip()
        )

    # --------------------------------------------------------
    # FUNDING INITIATIVE
    # --------------------------------------------------------

    if (
        "living lab" in lower
        and
        "water" in lower
    ):

        funding_initiative = (
            "Living Lab (Water)"
        )

    elif (
        "industrial water solutions"
        in lower
        or
        "wafer fab" in lower
    ):

        funding_initiative = (
            "Industrial Water Solutions (IWS)"
        )

    elif (
        "municipal water" in lower
        or
        "mwtd" in lower
    ):

        funding_initiative = (
            "Municipal Water: "
            "Technology Development (MWTD)"
        )

    elif (
        "competitive funding "
        "for water research"
        in lower
    ):

        funding_initiative = (
            "Competitive Funding "
            "for Water Research"
        )

    else:

        funding_initiative = None

    # --------------------------------------------------------
    # DOCUMENT TYPE
    # --------------------------------------------------------

    if (
        "funding initiative" in lower
        and
        "desired outcomes" in lower
    ):

        document_type = (
            "FI / programme paper"
        )

    elif (
        "project proposal" in lower
        or
        "scientific abstract" in lower
    ):

        document_type = (
            "Individual R&D proposal"
        )

    else:

        document_type = (
            "Unknown"
        )

    # --------------------------------------------------------
    # SECTION MAP
    # --------------------------------------------------------

    known_sections = [

        "Scientific Abstract",
        "Problem Statement",
        "Research Objectives",
        "Technical KPIs",
        "Methodology",
        "Landscape Scan",
        "Innovativeness",
        "Commercialisation",
        "Milestones",
        "Budget",
        "Impact",
        "TRL",

    ]

    sections = []

    for section in known_sections:

        if section.lower() in lower:

            sections.append({

                "name":
                    section,

                "confidence":
                    0.50,

            })

    # --------------------------------------------------------
    # POTENTIAL TECHNICAL CLAIMS
    # --------------------------------------------------------

    claims = []

    for line in text.splitlines():

        line = line.strip()

        if len(line) < 40:

            continue

        if re.search(

            r"novel|innovative|"
            r"improv|increase|"
            r"reduce|demonstrat|"
            r"achiev|target|"
            r"performance",

            line,

            re.IGNORECASE,

        ):

            claims.append(
                line
            )

        if len(claims) >= 12:

            break

    # --------------------------------------------------------
    # SIMPLE SUMMARY
    # --------------------------------------------------------

    meaningful_lines = [

        line.strip()

        for line in text.splitlines()

        if len(
            line.strip()
        ) > 80

    ]

    summary = None

    if meaningful_lines:

        summary = (
            " ".join(
                meaningful_lines[:3]
            )
        )

        if len(summary) > 800:

            summary = (
                summary[:800]
                + "..."
            )

    return {

        "document": {

            "filename":
                parsed.get(
                    "filename"
                ),

            "title":
                title,

            "funding_initiative":
                funding_initiative,

            "document_type":
                document_type,

            "pages":
                parsed.get(
                    "pages"
                ),

            "visual_pages":
                parsed.get(
                    "visual_pages",
                    []
                ),

            "sections":
                sections,

        },

        "understanding": {

            "problem":
                None,

            "technology":
                None,

            "baseline":
                None,

            "proposed_solution":
                None,

            "novelty_claims":
                [],

            "trl_start":
                None,

            "trl_target":
                None,

        },

        "claims":
            claims,

        "kpis":
            [],

        "summary":
            summary,

        "raw_text":
            parsed.get(
                "text",
                ""
            ),

        "review_flags": [

            {

                "severity":
                    "Review",

                "title":
                    "Fast-pass interpretation",

                "detail":
                    (
                        "The proposal has been "
                        "parsed quickly. "
                        "The deeper AI stage is "
                        "optional."
                    ),

            }

        ],

    }


# ============================================================
# RESEARCH SEARCH
# ============================================================

SEARCH_CACHE = {}


async def get_json(
    url: str,
) -> dict:

    if url in SEARCH_CACHE:

        return SEARCH_CACHE[
            url
        ]

    headers = {

        "Accept":
            "application/json",

        "User-Agent":
            "FI-Research-Intelligence/0.1",

    }

    async with httpx.AsyncClient(

        timeout=30,

        follow_redirects=True,

        headers=headers,

    ) as client:

        for attempt in range(4):

            response = await client.get(
                url
            )

            if response.status_code == 200:

                data = (
                    response.json()
                )

                SEARCH_CACHE[
                    url
                ] = data

                return data

            if (
                response.status_code
                == 429
            ):

                retry_after = (
                    response
                    .headers
                    .get(
                        "Retry-After"
                    )
                )

                try:

                    delay = (

                        float(
                            retry_after
                        )

                        if retry_after

                        else

                        1.5
                        * (
                            2 ** attempt
                        )

                    )

                except ValueError:

                    delay = (
                        1.5
                        * (
                            2 ** attempt
                        )
                    )

                await asyncio.sleep(
                    delay
                )

                continue

            response.raise_for_status()

    raise RuntimeError(
        "Research provider unavailable after retries."
    )


async def research_search(
    query: str,
    year: int,
    limit: int,
) -> list:

    query = query.strip()

    if not query:

        raise ValueError(
            "Research query cannot be empty."
        )

    # --------------------------------------------------------
    # OPENALEX
    # --------------------------------------------------------

    openalex_url = (

        "https://api.openalex.org/works"

        f"?search={quote(query)}"

        f"&filter="
        f"from_publication_date:{year}-01-01,"
        f"type:article|review"

        f"&per-page={limit}"

    )

    openalex_data = await get_json(
        openalex_url
    )

    openalex_results = []

    for item in openalex_data.get(
        "results",
        []
    ):

        location = (
            item.get(
                "primary_location"
            )
            or {}
        )

        openalex_results.append({

            "source":
                "OpenAlex",

            "title":
                item.get(
                    "title"
                )
                or
                "Untitled",

            "year":
                item.get(
                    "publication_year"
                ),

            "citations":
                item.get(
                    "cited_by_count",
                    0,
                ),

            "doi":
                item.get(
                    "doi"
                ),

            "url":
                location.get(
                    "landing_page_url"
                )
                or
                item.get("doi")
                or
                item.get("id"),

        })

    # --------------------------------------------------------
    # COURTESY DELAY
    # --------------------------------------------------------

    await asyncio.sleep(
        0.75
    )

    # --------------------------------------------------------
    # CROSSREF
    # --------------------------------------------------------

    crossref_url = (

        "https://api.crossref.org/works"

        f"?query.bibliographic="
        f"{quote(query)}"

        f"&filter="
        f"from-pub-date:{year}-01-01"

        f"&rows={limit}"

    )

    crossref_data = await get_json(
        crossref_url
    )

    crossref_results = []

    for item in (

        crossref_data

        .get(
            "message",
            {}
        )

        .get(
            "items",
            []
        )

    ):

        doi = item.get(
            "DOI"
        )

        date_parts = (

            item

            .get(
                "published",
                {}
            )

            .get(
                "date-parts",
                [[None]]
            )

        )

        crossref_results.append({

            "source":
                "Crossref",

            "title":
                (
                    item.get(
                        "title"
                    )
                    or
                    ["Untitled"]
                )[0],

            "year":
                (
                    date_parts[0][0]
                    if date_parts
                    else None
                ),

            "citations":
                item.get(
                    "is-referenced-by-count",
                    0,
                ),

            "doi":
                doi,

            "url":
                item.get(
                    "URL"
                )
                or
                (
                    f"https://doi.org/{doi}"
                    if doi
                    else None
                ),

        })

    # --------------------------------------------------------
    # COMBINE + DEDUP
    # --------------------------------------------------------

    combined = (
        openalex_results
        +
        crossref_results
    )

    seen = set()
    unique = []

    for item in combined:

        key = (

            item.get(
                "doi"
            )

            or

            item.get(
                "url"
            )

            or

            item.get(
                "title"
            )

            or
            ""

        ).lower()

        if (
            not key
            or key in seen
        ):

            continue

        seen.add(
            key
        )

        unique.append(
            item
        )

    unique.sort(

        key=lambda item:
            item.get(
                "citations",
                0,
            ),

        reverse=True,

    )

    return unique


# ============================================================
# PATENT / RESEARCH LINKS
# ============================================================

def direct_research_links(
    query: str,
) -> list:

    encoded = quote(
        query
    )

    return [

        (
            "Google Scholar",

            (
                "https://scholar.google.com/"
                f"scholar?q={encoded}"
            ),

        ),

        (
            "Google Patents",

            (
                "https://patents.google.com/"
                f"?q={encoded}"
            ),

        ),

        (
            "Semantic Scholar",

            (
                "https://www.semanticscholar.org/"
                f"search?q={encoded}"
            ),

        ),

        (
            "WIPO PATENTSCOPE",

            (
                "https://patentscope.wipo.int/"
                "search/en/result.jsf"
                f"?query={encoded}"
            ),

        ),

        (
            "USPTO",

            (
                "https://ppubs.uspto.gov/"
                "pubwebapp/static/pages/"
                "landing.html"
            ),

        ),

    ]


# ============================================================
# MAIN APPLICATION
# ============================================================

def main():

    state = {

        "proposal":
            None,

        "awards":
            [],

        "research":
            [],

        "award_results":
            [],

        "ai":
            None,

    }

    # ========================================================
    # OVERVIEW
    # ========================================================

    @ui.refreshable
    def overview():

        with ui.column().classes(

            "fi-shell "
            "w-full "
            "p-6 "
            "gap-6"

        ):

            # ------------------------------------------------
            # HERO
            # ------------------------------------------------

            with ui.card().classes(

                "fi-card "
                "w-full "
                "p-8 "
                "bg-white"

            ):

                ui.label(
                    "Research with context."
                ).classes(

                    "text-4xl "
                    "font-semibold "
                    "tracking-tight"

                )

                ui.label(

                    "Understand proposals, "
                    "find prior work, "
                    "and surface "
                    "the evidence that matters."

                ).classes(

                    "text-lg "
                    "text-slate-500 "
                    "max-w-3xl"

                )

            # ------------------------------------------------
            # METRICS
            # ------------------------------------------------

            with ui.grid(
                columns=4
            ).classes(
                "w-full gap-4"
            ):

                cards = [

                    (
                        "Proposal",

                        (
                            "Loaded"
                            if state["proposal"]
                            else
                            "None"
                        ),

                        "description",

                        "primary",

                    ),

                    (
                        "Awards",

                        len(
                            state[
                                "awards"
                            ]
                        ),

                        "workspace_premium",

                        "warning",

                    ),

                    (
                        "Research",

                        len(
                            state[
                                "research"
                            ]
                        ),

                        "science",

                        "positive",

                    ),

                    (
                        "AI",

                        (
                            "Ready"
                            if state["ai"]
                            else
                            "Optional"
                        ),

                        "auto_awesome",

                        "secondary",

                    ),

                ]

                for (
                    label,
                    value,
                    icon,
                    colour,
                ) in cards:

                    with ui.card().classes(

                        "fi-card "
                        "p-5 "
                        "bg-white"

                    ):

                        with ui.row().classes(

                            "w-full "
                            "items-center "
                            "justify-between"

                        ):

                            ui.label(
                                label
                            ).classes(
                                "text-sm "
                                "text-slate-500"
                            )

                            ui.icon(
                                icon,
                                color=colour,
                            )

                        ui.label(
                            str(value)
                        ).classes(

                            "text-2xl "
                            "font-semibold "
                            "mt-2"

                        )

            # ------------------------------------------------
            # PROPOSAL
            # ------------------------------------------------

            if not state["proposal"]:

                with ui.card().classes(

                    "fi-card "
                    "w-full "
                    "p-6 "

                ):

                    ui.icon(

                        "description",

                        size="2.8rem",

                        color="sky-300",

                    )

                    ui.label(
                        "No proposal loaded"
                    ).classes(
                        "text-xl font-semibold"
                    )

                    ui.label(

                        "Go to Document and "
                        "upload a proposal."

                    ).classes(
                        "text-slate-500"
                    )

                return

            proposal = state[
                "proposal"
            ]

            document = proposal[
                "document"
            ]

            understanding = proposal[
                "understanding"
            ]

            with ui.card().classes(

                "fi-card "
                "w-full "
                "p-6 "

            ):

                ui.label(

                    document.get(
                        "title"
                    )

                    or

                    "Title not confidently determined"

                ).classes(

                    "text-2xl "
                    "font-semibold"

                )

                with ui.row().classes(
                    "gap-2 mt-2"
                ):

                    ui.chip(

                        document.get(
                            "funding_initiative"
                        )
                        or
                        "FI not determined",

                        icon="account_balance",

                        color="primary",

                    )

                    ui.chip(

                        document.get(
                            "document_type"
                        )
                        or
                        "Unknown",

                        icon="description",

                    )

                ui.separator().classes(
                    "my-4"
                )

                with ui.grid(
                    columns=2
                ).classes(
                    "w-full gap-6"
                ):

                    with ui.column():

                        ui.label(
                            "Summary"
                        ).classes(

                            "text-xs "
                            "uppercase "
                            "font-semibold "
                            "text-slate-400"

                        )

                        ui.label(

                            proposal.get(
                                "summary"
                            )
                            or
                            "No fast-pass summary available."

                        ).classes(
                            "text-slate-700"
                        )

                    with ui.column():

                        ui.label(
                            "Pages flagged for visual review"
                        ).classes(

                            "text-xs "
                            "uppercase "
                            "font-semibold "
                            "text-slate-400"

                        )

                        ui.label(

                            str(
                                document.get(
                                    "visual_pages",
                                    []
                                )
                            )

                        ).classes(
                            "text-slate-700"
                        )

                if state["ai"]:

                    ui.separator().classes(
                        "my-5"
                    )

                    ui.label(
                        "AI findings"
                    ).classes(
                        "text-xl font-semibold"
                    )

                    ai = state[
                        "ai"
                    ]

                    with ui.grid(
                        columns=2
                    ).classes(
                        "w-full gap-4 mt-3"
                    ):

                        for (
                            label,
                            key,
                        ) in [

                            (
                                "Problem",
                                "problem",
                            ),

                            (
                                "Technology",
                                "technology",
                            ),

                            (
                                "Baseline",
                                "baseline",
                            ),

                            (
                                "Proposed solution",
                                "proposed_solution",
                            ),

                            (
                                "TRL",
                                "trl",
                            ),

                            (
                                "Commercialisation",
                                "commercialisation",
                            ),

                        ]:

                            with ui.card().classes(
                                "fi-card "
                                "p-4 "
                                "bg-sky-50"
                            ):

                                ui.label(
                                    label
                                ).classes(
                                    "text-xs "
                                    "uppercase "
                                    "font-semibold "
                                    "text-slate-400"
                                )

                                ui.label(

                                    ai.get(
                                        key
                                    )
                                    or
                                    "Not determined"

                                ).classes(
                                    "text-slate-700"
                                )

                ui.button(

                    "Run deeper AI analysis",

                    icon="auto_awesome",

                    on_click=lambda:
                        asyncio.create_task(
                            run_ai()
                        ),

                ).props(

                    "color=primary "
                    "unelevated "
                    "no-caps"

                ).classes(
                    "mt-5"
                )


    # ========================================================
    # DOCUMENT VIEW
    # ========================================================

    @ui.refreshable
    def document_panel():

        with ui.column().classes(

            "fi-shell "
            "w-full "
            "p-6 "
            "gap-5"

        ):

            ui.label(
                "Document understanding"
            ).classes(
                "text-3xl font-semibold"
            )

            ui.label(

                "Native extraction first. "
                "AI interpretation second."

            ).classes(
                "text-slate-500"
            )

            proposal = state[
                "proposal"
            ]

            if not proposal:

                ui.label(
                    "No proposal loaded."
                ).classes(
                    "text-slate-400"
                )

                return

            d = proposal[
                "document"
            ]

            u = proposal[
                "understanding"
            ]

            with ui.grid(
                columns=2
            ).classes(
                "w-full gap-4"
            ):

                fields = [

                    (
                        "Problem",
                        "problem",
                    ),

                    (
                        "Technology",
                        "technology",
                    ),

                    (
                        "Baseline",
                        "baseline",
                    ),

                    (
                        "Proposed solution",
                        "proposed_solution",
                    ),

                ]

                for (
                    label,
                    key,
                ) in fields:

                    with ui.card().classes(
                        "fi-card "
                        "w-full "
                        "p-5"
                    ):

                        ui.label(
                            label
                        ).classes(
                            "font-semibold"
                        )

                        ui.label(

                            u.get(
                                key
                            )
                            or
                            "Not determined"

                        ).classes(
                            "text-slate-600 "
                            "mt-2"
                        )

            with ui.card().classes(
                "fi-card "
                "w-full "
                "p-5"
            ):

                ui.label(
                    "Document map"
                ).classes(
                    "text-xl font-semibold"
                )

                rows = [

                    {

                        "section":
                            item.get(
                                "name"
                            ),

                        "confidence":
                            (
                                f"{round("
                                f"(item.get('confidence') or 0)"
                                f" * 100)}%"
                            ),

                    }

                    for item in d.get(
                        "sections",
                        []
                    )

                ]

                ui.table(

                    columns=[

                        {

                            "name":
                                "section",

                            "label":
                                "SECTION",

                            "field":
                                "section",

                        },

                        {

                            "name":
                                "confidence",

                            "label":
                                "CONFIDENCE",

                            "field":
                                "confidence",

                        },

                    ],

                    rows=rows,

                ).props(
                    "flat bordered"
                ).classes(
                    "w-full"
                )

            with ui.card().classes(
                "fi-card "
                "w-full "
                "p-5"
            ):

                ui.label(
                    "Potential technical claims"
                ).classes(
                    "text-xl font-semibold"
                )

                claims = proposal.get(
                    "claims",
                    []
                )

                if not claims:

                    ui.label(
                        "No claim signals detected."
                    ).classes(
                        "text-slate-400"
                    )

                for claim in claims:

                    ui.label(
                        "• " + str(claim)
                    ).classes(
                        "text-slate-700"
                    )


    # ========================================================
    # AWARDS
    # ========================================================

    @ui.refreshable
    def awards_panel():

        with ui.column().classes(

            "fi-shell "
            "w-full "
            "p-6 "
            "gap-5"

        ):

            ui.label(
                "Award landscape"
            ).classes(
                "text-3xl font-semibold"
            )

            ui.label(

                "Load previous awards "
                "or proposals for comparison."

            ).classes(
                "text-slate-500"
            )

            ui.upload(

                on_upload=lambda event:
                    asyncio.create_task(
                        handle_award(event)
                    ),

                multiple=True,

                auto_upload=True,

                max_files=100,

                max_file_size=50_000_000,

            ).props(

                "accept=.pdf,.docx,.txt,.md"

            ).classes(
                "w-full"
            )

            if not state[
                "awards"
            ]:

                ui.label(
                    "No awards loaded."
                ).classes(
                    "text-slate-400"
                )

            else:

                for award in state[
                    "awards"
                ]:

                    with ui.card().classes(
                        "fi-card "
                        "w-full "
                        "p-4"
                    ):

                        ui.label(
                            award[
                                "filename"
                            ]
                        ).classes(
                            "font-semibold"
                        )

                if state[
                    "proposal"
                ]:

                    ui.button(

                        "Compare awards",

                        icon="compare_arrows",

                        on_click=lambda:
                            asyncio.create_task(
                                compare_awards()
                            ),

                    ).props(

                        "color=primary "
                        "unelevated "
                        "no-caps"

                    )

                for item in state[
                    "award_results"
                ]:

                    with ui.card().classes(
                        "fi-card "
                        "w-full "
                        "p-4 "
                        "bg-slate-50"
                    ):

                        with ui.row().classes(
                            "items-center"
                        ):

                            ui.label(
                                item[
                                    "filename"
                                ]
                            ).classes(
                                "font-semibold"
                            )

                            ui.space()

                            ui.badge(
                                item[
                                    "relationship"
                                ],
                                color="warning",
                            )

                        ui.label(
                            item[
                                "reason"
                            ]
                        ).classes(
                            "text-slate-600 "
                            "mt-2"
                        )


    # ========================================================
    # RESEARCH
    # ========================================================

    @ui.refreshable
    def research_panel():

        with ui.column().classes(

            "fi-shell "
            "w-full "
            "p-6 "
            "gap-5"

        ):

            ui.label(
                "Research & IP"
            ).classes(
                "text-3xl font-semibold"
            )

            ui.label(

                "Search literature "
                "and follow direct "
                "patent routes."

            ).classes(
                "text-slate-500"
            )

            query = ui.input(

                "Research question "
                "or technical claim"

            ).classes(
                "w-full"
            )

            year = ui.number(

                "From year",

                value=2020,

                min=1900,

                max=2100,

            )

            limit = ui.number(

                "Results",

                value=10,

                min=1,

                max=25,

            )

            async def search():

                note = ui.notification(

                    "Searching OpenAlex and Crossref...",

                    spinner=True,

                    timeout=None,

                )

                try:

                    state[
                        "research"
                    ] = await research_search(

                        query.value or "",

                        int(
                            year.value
                        ),

                        int(
                            limit.value
                        ),

                    )

                    note.message = (

                        f"Found "
                        f"{len(state['research'])} "
                        f"records."

                    )

                    note.spinner = False

                    note.type = (
                        "positive"
                    )

                    research_panel.refresh()
                    overview.refresh()

                except Exception as error:

                    note.message = (
                        f"Search failed: "
                        f"{error}"
                    )

                    note.spinner = False

                    note.type = (
                        "negative"
                    )

            ui.button(

                "Search",

                icon="search",

                on_click=search,

            ).props(

                "color=primary "
                "unelevated "
                "no-caps"

            )

            if state["proposal"]:

                proposal_title = (

                    state["proposal"]

                    .get("document", {})

                    .get("title")

                )

                if (
                    proposal_title
                    and not query.value
                ):

                    query.value = (
                        proposal_title
                    )

            link_query = (

                query.value

                or

                (
                    (
                        state["proposal"]
                        or {}
                    )
                    .get(
                        "document",
                        {}
                    )
                    .get(
                        "title"
                    )
                )

                or

                "water technology"

            )

            with ui.row().classes(
                "gap-5 flex-wrap"
            ):

                for (
                    name,
                    url,
                ) in direct_research_links(
                    link_query
                ):

                    ui.link(
                        name,
                        url,
                        new_tab=True,
                    ).classes(
                        "text-sky-600 "
                        "font-semibold"
                    )

            if not state[
                "research"
            ]:

                ui.label(
                    "No research results yet."
                ).classes(
                    "text-slate-400"
                )

            else:

                for paper in state[
                    "research"
                ]:

                    with ui.card().classes(
                        "fi-card "
                        "w-full "
                        "p-5"
                    ):

                        with ui.row().classes(
                            "items-center "
                            "w-full"
                        ):

                            ui.badge(

                                paper[
                                    "source"
                                ],

                                color=(

                                    "positive"

                                    if paper[
                                        "source"
                                    ]
                                    ==
                                    "OpenAlex"

                                    else

                                    "warning"

                                ),

                            )

                            ui.label(

                                str(
                                    paper.get(
                                        "year"
                                    )
                                    or
                                    ""
                                )

                            ).classes(
                                "text-slate-400"
                            )

                            ui.space()

                            ui.label(

                                f"{paper.get('citations', 0)} citations"

                            ).classes(
                                "text-slate-400"
                            )

                        ui.label(

                            paper.get(
                                "title"
                            )
                            or
                            "Untitled"

                        ).classes(

                            "font-semibold "
                            "text-slate-800"

                        )

                        if paper.get(
                            "url"
                        ):

                            ui.link(

                                "Open source ↗",

                                paper[
                                    "url"
                                ],

                                new_tab=True,

                            ).classes(

                                "text-sky-600 "
                                "font-semibold"

                            )


    # ========================================================
    # REVIEWER
    # ========================================================

    @ui.refreshable
    def reviewer_panel():

        with ui.column().classes(

            "fi-shell "
            "w-full "
            "p-6 "
            "gap-5"

        ):

            ui.label(
                "Human reviewer"
            ).classes(
                "text-3xl font-semibold"
            )

            ui.label(

                "The system provides evidence. "
                "The panel makes the final decision."

            ).classes(
                "text-slate-500"
            )

            criteria = [

                "Science & technology",

                "Impact / national benefit",

                "Management & delivery",

                "Budget / value",

            ]

            sliders = []

            with ui.grid(
                columns=2
            ).classes(
                "w-full gap-4"
            ):

                for criterion in criteria:

                    with ui.card().classes(
                        "fi-card "
                        "p-5"
                    ):

                        ui.label(
                            criterion
                        ).classes(
                            "font-semibold"
                        )

                        slider = ui.slider(

                            min=1,

                            max=5,

                            step=1,

                            value=4,

                        ).classes(
                            "w-full"
                        )

                        ui.label().bind_text_from(

                            slider,

                            "value",

                            backward=lambda value:

                                f"Score: "
                                f"{value}/5",

                        )

                        sliders.append(
                            slider
                        )

            overall = ui.label(
                "Overall: 4.00 / 5"
            ).classes(

                "text-3xl "
                "font-semibold "
                "text-sky-600"

            )

            def update_score():

                values = [

                    float(
                        slider.value
                        or 0
                    )

                    for slider
                    in sliders

                ]

                average = (

                    sum(values)
                    /
                    len(values)

                    if values

                    else

                    0

                )

                overall.set_text(

                    f"Overall: "
                    f"{average:.2f} / 5"

                )

            for slider in sliders:

                slider.on_value_change(
                    update_score
                )


    # ========================================================
    # PROPOSAL UPLOAD
    # ========================================================

    async def handle_proposal(
        event
    ):

        note = ui.notification(

            "Receiving proposal...",

            spinner=True,

            timeout=None,

        )

        path = None

        try:

            suffix = Path(
                event.name
            ).suffix.lower()

            content = await run.io_bound(

                event.content.read

            )

            with tempfile.NamedTemporaryFile(

                delete=False,

                suffix=suffix,

            ) as temp:

                temp.write(
                    content
                )

                path = (
                    temp.name
                )

            note.message = (
                "Extracting document text..."
            )

            parsed = await run.io_bound(

                parse_document,

                path,

                event.name,

            )

            parsed[
                "filename"
            ] = event.name

            state[
                "proposal"
            ] = await run.io_bound(

                understand_document,

                parsed,

            )

            state[
                "ai"
            ] = None

            note.message = (
                "Proposal ready."
            )

            note.spinner = False

            note.type = (
                "positive"
            )

            overview.refresh()

            document_panel.refresh()

            research_panel.refresh()

        except Exception as error:

            note.message = (
                f"Proposal upload failed: "
                f"{error}"
            )

            note.spinner = False

            note.type = (
                "negative"
            )

        finally:

            if path:

                try:

                    os.unlink(
                        path
                    )

                except OSError:

                    pass


    # ========================================================
    # AWARD UPLOAD
    # ========================================================

    async def handle_award(
        event
    ):

        note = ui.notification(

            f"Reading {event.name}...",

            spinner=True,

            timeout=None,

        )

        path = None

        try:

            suffix = Path(
                event.name
            ).suffix.lower()

            content = await run.io_bound(

                event.content.read

            )

            with tempfile.NamedTemporaryFile(

                delete=False,

                suffix=suffix,

            ) as temp:

                temp.write(
                    content
                )

                path = (
                    temp.name
                )

            parsed = await run.io_bound(

                parse_document,

                path,

                event.name,

            )

            state[
                "awards"
            ].append({

                "filename":
                    event.name,

                "text":
                    parsed[
                        "text"
                    ],

            })

            note.message = (
                f"{event.name} loaded."
            )

            note.spinner = False

            note.type = (
                "positive"
            )

            awards_panel.refresh()

            overview.refresh()

        except Exception as error:

            note.message = (
                f"Award upload failed: "
                f"{error}"
            )

            note.spinner = False

            note.type = (
                "negative"
            )

        finally:

            if path:

                try:

                    os.unlink(
                        path
                    )

                except OSError:

                    pass


    # ========================================================
    # DEEPER AI
    # ========================================================

    async def run_ai():

        proposal = state[
            "proposal"
        ]

        if not proposal:

            ui.notify(

                "Upload a proposal first.",

                type="warning",

            )

            return

        try:

            # Import only when requested.
            from openai import AsyncOpenAI

            api_key = os.getenv(
                "OPENAI_API_KEY"
            )

            if not api_key:

                raise RuntimeError(

                    "OPENAI_API_KEY is not configured. "
                    "The free features still work."

                )

            client = AsyncOpenAI(
                api_key=api_key
            )

            system_prompt = """

You are a senior R&D funding proposal analyst.

Read the supplied proposal carefully.

Do not invent information.

Separate applicant claims from established facts.

Use null where unsupported.

Return JSON only with:

confidence
problem
technology
baseline
proposed_solution
trl
commercialisation
novelty_claims
prior_projects
research_questions
review_flags

Each review_flags item should contain:

title
detail
severity

"""

            payload = {

                "document":
                    proposal.get(
                        "document",
                        {}
                    ),

                "understanding":
                    proposal.get(
                        "understanding",
                        {}
                    ),

                "claims":
                    proposal.get(
                        "claims",
                        []
                    ),

                "source_text":
                    proposal.get(
                        "raw_text",
                        ""
                    )[:180000],

            }

            response = await client.responses.create(

                model=os.getenv(
                    "OPENAI_MODEL",
                    "gpt-5",
                ),

                input=[

                    {

                        "role":
                            "system",

                        "content": [

                            {

                                "type":
                                    "input_text",

                                "text":
                                    system_prompt,

                            }

                        ],

                    },

                    {

                        "role":
                            "user",

                        "content": [

                            {

                                "type":
                                    "input_text",

                                "text":
                                    __import__(
                                        "json"
                                    ).dumps(
                                        payload,
                                        ensure_ascii=False,
                                    ),

                            }

                        ],

                    },

                ],

                text={

                    "format": {

                        "type":
                            "json_object",

                    }

                },

            )

            state[
                "ai"
            ] = __import__(
                "json"
            ).loads(
                response.output_text
            )

            overview.refresh()

            ui.notify(

                "AI analysis complete.",

                type="positive",

            )

        except Exception as error:

            ui.notify(

                str(error),

                type="negative",

            )


    # ========================================================
    # AWARD COMPARISON
    # ========================================================

    async def compare_awards():

        proposal = state[
            "proposal"
        ]

        if not proposal:

            ui.notify(

                "Upload a current proposal first.",

                type="warning",

            )

            return

        proposal_text = (

            proposal.get(
                "raw_text",
                ""
            )
            .lower()

        )

        terms = [

            "membrane",
            "sludge",
            "wastewater",
            "desalination",
            "pfas",
            "carbon",
            "resource recovery",
            "anaerobic",
            "electrolysis",
            "biogas",
            "water reuse",

        ]

        results = []

        for award in state[
            "awards"
        ]:

            text = award[
                "text"
            ].lower()

            shared = [

                term

                for term in terms

                if term in proposal_text
                and term in text

            ]

            if len(shared) >= 4:

                relationship = (
                    "Potential overlap"
                )

            elif len(shared) >= 2:

                relationship = (
                    "Related / complementary"
                )

            elif len(shared) == 1:

                relationship = (
                    "Weakly related"
                )

            else:

                relationship = (
                    "No strong relationship found"
                )

            results.append({

                "filename":
                    award[
                        "filename"
                    ],

                "relationship":
                    relationship,

                "reason":
                    (
                        "Shared signals: "
                        +
                        (
                            ", ".join(shared)
                            if shared
                            else
                            "none detected"
                        )
                    ),

            })

        state[
            "award_results"
        ] = results

        awards_panel.refresh()


    # ========================================================
    # HEADER
    # ========================================================

    with ui.header().classes(

        "bg-white "
        "text-slate-900 "
        "border-b border-sky-100 "
        "px-6 py-3"

    ):

        with ui.row().classes(

            "w-full "
            "fi-shell "
            "items-center "
            "gap-3"

        ):

            with ui.element(
                "div"
            ).classes(

                "w-9 h-9 "
                "rounded-xl "
                "bg-sky-500 "
                "flex items-center "
                "justify-center"

            ):

                ui.icon(
                    "water",
                    color="white",
                )

            with ui.column().classes(
                "gap-0"
            ):

                ui.label(
                    APP_TITLE
                ).classes(

                    "text-base "
                    "font-semibold"

                )

                ui.label(
                    "Proposal research workspace"
                ).classes(

                    "text-xs "
                    "text-slate-400"

                )

            ui.space()

            ui.badge(
                "BETA",
                color="primary",
            )


    # ========================================================
    # TABS
    # ========================================================

    with ui.tabs() as tabs:

        ui.tab(
            "Overview",
            icon="dashboard",
        )

        ui.tab(
            "Document",
            icon="description",
        )

        ui.tab(
            "Awards",
            icon="workspace_premium",
        )

        ui.tab(
            "Research & IP",
            icon="science",
        )

        ui.tab(
            "Reviewer",
            icon="fact_check",
        )


    with ui.tab_panels(

        tabs,

        value="Overview",

    ).classes(

        "w-full "
        "bg-transparent"

    ):

        with ui.tab_panel(
            "Overview"
        ):

            overview()


        with ui.tab_panel(
            "Document"
        ):

            document_panel()


        with ui.tab_panel(
            "Awards"
        ):

            awards_panel()


        with ui.tab_panel(
            "Research & IP"
        ):

            research_panel()


        with ui.tab_panel(
            "Reviewer"
        ):

            reviewer_panel()


# ============================================================
# START
# ============================================================

if __name__ == "__main__":

    ui.run(

        host="0.0.0.0",

        port=int(
            os.environ.get(
                "PORT",
                "10000",
            )
        ),

        title=APP_TITLE,

        favicon="🔬",

        reload=False,

        show=False,

    )
