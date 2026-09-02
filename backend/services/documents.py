from pathlib import Path
import re


SECTION_NAMES = [
    "scientific abstract",
    "abstract",
    "executive summary",
    "problem statement",
    "background",
    "research objectives",
    "objectives",
    "technical kpis",
    "key performance indicators",
    "methodology",
    "approach",
    "landscape scan",
    "literature review",
    "innovativeness",
    "innovation",
    "commercialisation",
    "commercialization",
    "milestones",
    "budget",
    "impact",
    "trl",
]


TECHNICAL_PATTERNS = [
    r"\bceramic membranes?\b",
    r"\bmembranes?\b",
    r"\bdesalination\b",
    r"\bwastewater\b",
    r"\bwater treatment\b",
    r"\bwater reuse\b",
    r"\breverse osmosis\b",
    r"\bultrafiltration\b",
    r"\bmicrofiltration\b",
    r"\bnanofiltration\b",
    r"\banaerobic\b",
    r"\belectrolysis\b",
    r"\belectrochemical\b",
    r"\badsorption\b",
    r"\boxidation\b",
    r"\bfiltration\b",
    r"\bresource recovery\b",
    r"\bcarbon capture\b",
    r"\bpfas\b",
    r"\bsludge\b",
    r"\bbiogas\b",
]


CLAIM_PATTERN = re.compile(
    r"\b("
    r"novel|"
    r"innovative|"
    r"improve|"
    r"improved|"
    r"improvement|"
    r"increase|"
    r"increased|"
    r"reduction|"
    r"reduce|"
    r"reduced|"
    r"achieve|"
    r"achieved|"
    r"target|"
    r"demonstrate|"
    r"demonstrated|"
    r"performance|"
    r"efficiency|"
    r"higher|"
    r"lower|"
    r"enhance|"
    r"enhanced"
    r")\b",
    re.IGNORECASE,
)


KPI_PATTERN = re.compile(
    r"("
    r"\d+(?:\.\d+)?\s*%|"
    r"\d+(?:\.\d+)?\s*"
    r"(?:"
    r"mg\/l|"
    r"g\/l|"
    r"kg\/m3|"
    r"m3\/day|"
    r"m³\/day|"
    r"bar|"
    r"kwh|"
    r"kw|"
    r"mw|"
    r"ppm|"
    r"ppb"
    r")|"
    r"trl\s*\d+"
    r")",
    re.IGNORECASE,
)


def parse_document(
    path: str,
    filename: str,
) -> dict:

    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return parse_pdf(
            path,
            filename,
        )

    if suffix == ".docx":
        return parse_docx(
            path,
            filename,
        )

    if suffix in {
        ".txt",
        ".md",
    }:
        return parse_text(
            path,
            filename,
        )

    raise ValueError(
        "Unsupported document type. "
        "Please upload PDF, DOCX, TXT or MD."
    )


def parse_pdf(
    path: str,
    filename: str,
) -> dict:

    import pymupdf

    document = pymupdf.open(path)

    pages = []
    total_chars = 0

    try:

        for page_number, page in enumerate(
            document,
            start=1,
        ):

            # ------------------------------------------------
            # First extraction mode
            # ------------------------------------------------

            text = (
                page.get_text(
                    "text",
                    sort=True,
                )
                or ""
            ).strip()

            # ------------------------------------------------
            # Fallback extraction
            # ------------------------------------------------

            if len(text) < 20:

                blocks = (
                    page.get_text(
                        "blocks"
                    )
                    or []
                )

                block_text = []

                for block in blocks:

                    if len(block) >= 5:

                        value = str(
                            block[4]
                        ).strip()

                        if value:

                            block_text.append(
                                value
                            )

                fallback_text = "\n".join(
                    block_text
                ).strip()

                if len(fallback_text) > len(text):

                    text = fallback_text

            text_length = len(text)

            total_chars += text_length

            pages.append(
                {
                    "page": page_number,
                    "text": text,
                    "text_length": text_length,
                    "needs_visual_review": (
                        text_length < 200
                    ),
                    "has_images": bool(
                        page.get_images(
                            full=True
                        )
                    ),
                    "image_count": len(
                        page.get_images(
                            full=True
                        )
                    ),
                }
            )

    finally:

        document.close()

    if not pages:

        raise ValueError(
            "The PDF contains no pages."
        )

    scanned_pages = [
        page["page"]
        for page in pages
        if page["text_length"] < 200
    ]

    # If almost everything has no text layer,
    # the PDF is probably scanned/image-based.
    scanned_document = (
        total_chars < max(
            500,
            len(pages) * 100
        )
    )

    return {
        "filename": filename,
        "text": "\n\n".join(
            page["text"]
            for page in pages
            if page["text"]
        ),
        "pages": pages,
        "page_count": len(pages),
        "scanned_document": scanned_document,
        "scanned_pages": scanned_pages,
        "total_characters": total_chars,
    }


def parse_docx(
    path: str,
    filename: str,
) -> dict:

    from docx import Document

    document = Document(
        path
    )

    blocks = []

    for paragraph in document.paragraphs:

        text = (
            paragraph.text
            .strip()
        )

        if text:

            blocks.append(
                text
            )

    for table in document.tables:

        for row in table.rows:

            row_text = " | ".join(
                cell.text.strip()
                for cell in row.cells
            )

            if row_text:

                blocks.append(
                    row_text
                )

    text = "\n".join(
        blocks
    )

    return {
        "filename": filename,
        "text": text,
        "pages": [
            {
                "page": 1,
                "text": text,
                "text_length": len(text),
                "needs_visual_review": False,
                "has_images": False,
                "image_count": 0,
            }
        ],
        "page_count": None,
        "scanned_document": False,
        "scanned_pages": [],
        "total_characters": len(text),
    }


def parse_text(
    path: str,
    filename: str,
) -> dict:

    text = Path(
        path
    ).read_text(
        encoding="utf-8",
        errors="ignore",
    )

    return {
        "filename": filename,
        "text": text,
        "pages": [
            {
                "page": 1,
                "text": text,
                "text_length": len(text),
                "needs_visual_review": False,
                "has_images": False,
                "image_count": 0,
            }
        ],
        "page_count": 1,
        "scanned_document": False,
        "scanned_pages": [],
        "total_characters": len(text),
    }


def extract_title(
    text: str,
    filename: str,
) -> str:

    patterns = [
        (
            r"(?:"
            r"title of research project|"
            r"research project title|"
            r"proposal title|"
            r"project title"
            r")"
            r"\s*[:\-]?\s*"
            r"([^\n]{8,200})"
        ),
    ]

    for pattern in patterns:

        match = re.search(
            pattern,
            text,
            re.IGNORECASE,
        )

        if match:

            title = (
                match.group(
                    1
                )
                .strip()
            )

            if title:

                return title

    for line in text.splitlines():

        clean = (
            line
            .strip()
        )

        if (
            15 <= len(clean) <= 180
            and not clean.startswith("[")
        ):

            return clean

    return Path(
        filename
    ).stem


def extract_concepts(
    text: str,
) -> list:

    concepts = []

    for pattern in TECHNICAL_PATTERNS:

        matches = re.findall(
            pattern,
            text,
            re.IGNORECASE,
        )

        for match in matches:

            concept = (
                match
                if isinstance(
                    match,
                    str,
                )
                else match[0]
            )

            concept = (
                concept
                .strip()
                .lower()
            )

            if (
                concept
                and concept not in concepts
            ):

                concepts.append(
                    concept
                )

    return concepts[:15]


def analyze_page(
    page: dict,
) -> dict:

    text = page.get(
        "text",
        "",
    )

    lower = text.lower()

    sections = []

    for section in SECTION_NAMES:

        if section in lower:

            display_name = (
                section.title()
            )

            if (
                display_name
                not in sections
            ):

                sections.append(
                    display_name
                )

    concepts = extract_concepts(
        text
    )

    claims = []
    kpis = []

    for line in text.splitlines():

        clean = (
            line
            .strip()
        )

        if len(clean) < 30:

            continue

        if CLAIM_PATTERN.search(
            clean
        ):

            if clean not in claims:

                claims.append(
                    clean
                )

        if KPI_PATTERN.search(
            clean
        ):

            if clean not in kpis:

                kpis.append(
                    clean
                )

        if (
            len(claims) >= 8
            and len(kpis) >= 8
        ):

            break

    return {
        "page": page.get(
            "page"
        ),
        "sections": sections[:8],
        "concepts": concepts[:12],
        "claims": claims[:8],
        "kpis": kpis[:8],
        "needs_visual_review": page.get(
            "needs_visual_review",
            False,
        ),
        "has_images": page.get(
            "has_images",
            False,
        ),
        "image_count": page.get(
            "image_count",
            0,
        ),
        "text_preview": (
            text[:1500]
            if text
            else ""
        ),
    }


def build_document_dossier(
    parsed: dict,
) -> dict:

    text = parsed.get(
        "text",
        "",
    )

    filename = parsed.get(
        "filename",
        "proposal",
    )

    if not text.strip():

        raise ValueError(

            "No readable text was extracted "
            "from this document. This PDF may "
            "be scanned or image-only. "
            "Visual/OCR processing is required."

        )

    title = extract_title(
        text,
        filename,
    )

    page_analysis = [
        analyze_page(
            page
        )
        for page in parsed.get(
            "pages",
            [],
        )
    ]

    concepts = []
    claims = []
    kpis = []
    sections = []

    visual_review_pages = []

    for page in page_analysis:

        page_number = page.get(
            "page"
        )

        for concept in page.get(
            "concepts",
            [],
        ):

            if concept not in concepts:

                concepts.append(
                    concept
                )

        for claim in page.get(
            "claims",
            [],
        ):

            claims.append(
                {
                    "page":
                        page_number,

                    "text":
                        claim,
                }
            )

        for kpi in page.get(
            "kpis",
            [],
        ):

            kpis.append(
                {
                    "page":
                        page_number,

                    "text":
                        kpi,
                }
            )

        for section in page.get(
            "sections",
            [],
        ):

            sections.append(
                {
                    "page":
                        page_number,

                    "name":
                        section,
                }
            )

        if page.get(
            "needs_visual_review",
            False,
        ):

            visual_review_pages.append(
                page_number
            )

    unique_claims = []

    seen_claims = set()

    for claim in claims:

        key = (
            claim.get(
                "text",
                "",
            )
            .strip()
            .lower()
        )

        if (
            not key
            or key in seen_claims
        ):

            continue

        seen_claims.add(
            key
        )

        unique_claims.append(
            claim
        )

    unique_kpis = []

    seen_kpis = set()

    for kpi in kpis:

        key = (
            kpi.get(
                "text",
                "",
            )
            .strip()
            .lower()
        )

        if (
            not key
            or key in seen_kpis
        ):

            continue

        seen_kpis.add(
            key
        )

        unique_kpis.append(
            kpi
        )

    unique_sections = []

    seen_sections = set()

    for section in sections:

        key = (
            section.get(
                "page"
            ),
            section.get(
                "name"
            ),
        )

        if key in seen_sections:

            continue

        seen_sections.add(
            key
        )

        unique_sections.append(
            section
        )

    return {

        "document": {

            "filename":
                filename,

            "title":
                title,

            "pages":
                parsed.get(
                    "page_count"
                ),

            "total_characters":
                parsed.get(
                    "total_characters",
                    len(text),
                ),

            "scanned_document":
                parsed.get(
                    "scanned_document",
                    False,
                ),

            "scanned_pages":
                parsed.get(
                    "scanned_pages",
                    [],
                ),

            "visual_review_pages":
                visual_review_pages,

        },

        "concepts":
            concepts[:15],

        "claims":
            unique_claims[:30],

        "kpis":
            unique_kpis[:30],

        "sections":
            unique_sections,

        "page_analysis":
            page_analysis,

        "raw_text":
            text,

    }
